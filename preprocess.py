import streamlit as st
import os
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document as LangchainDocument
import requests
from bs4 import BeautifulSoup
import faiss
from dotenv import load_dotenv

load_dotenv()

def extract_text_from_pdf(file):
    paragraphs = []
    reader = PdfReader(file)
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            paragraphs.extend(page_text.split("\n"))
    return paragraphs

def extract_text_from_docx(file):
    paragraphs = []
    docx = DocxDocument(file)
    for paragraph in docx.paragraphs:
        if paragraph.text.strip():
            paragraphs.append(paragraph.text)
    return paragraphs

def extract_text_from_web(links):
    paragraphs = []
    for link in links:
        try:
            headers = {"User-Agent": "Mozilla/5.0"}
            response = requests.get(link, headers=headers, timeout=10)
            response.raise_for_status()
            soup = BeautifulSoup(response.text, "html.parser")
            
            for p in soup.find_all("p"):
                text = p.get_text(strip=True)
                if text:
                    paragraphs.append(text)
            
            faq_sections = soup.select(".faqs .accordion_toggle, .accordion_body")
            for faq in faq_sections:
                text = faq.get_text(strip=True)
                if text and text not in paragraphs:
                    paragraphs.append(text)
        except Exception as e:
            print(f"Failed to process link {link}: {e}")
    return paragraphs

def preprocess_text(files, links, size, overlap):
    paragraphs = []
    
    for file in files:
        if file.name.endswith(".pdf"):
            paragraphs.extend(extract_text_from_pdf(file))
        elif file.name.endswith(".docx"):
            paragraphs.extend(extract_text_from_docx(file))
    
    paragraphs.extend(extract_text_from_web(links))
    
    paragraphs = [para.strip() for para in paragraphs if para.strip()]
    docs = [LangchainDocument(page_content=para) for para in paragraphs]
    
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=size, chunk_overlap=overlap)
    text_chunks = text_splitter.split_documents(docs)
    
    return text_chunks


import os
import numpy as np
import sqlite3
import uuid
import streamlit as st
from langchain.embeddings import SentenceTransformerEmbeddings
from langchain.docstore.document import Document
from langchain.vectorstores import Chroma, FAISS
from langchain.docstore.in_memory import InMemoryDocstore
from qdrant_client import QdrantClient, models
from pinecone.grpc import PineconeGRPC as Pinecone
from pinecone import ServerlessSpec
import weaviate
from weaviate.auth import AuthApiKey
from weaviate.collections import Collection
from langchain_weaviate.vectorstores import WeaviateVectorStore

def preprocess_chroma(text, embedding_model, persist_directory):
    vectordb = Chroma.from_documents(documents=text, embedding=embedding_model, persist_directory=persist_directory)
    vectordb.persist()
    retriever = vectordb.as_retriever()
    return vectordb, retriever

def preprocess_faiss(text, embedding_model):
    texts = [doc.page_content for doc in text]
    embeddings = np.array(embedding_model.embed_documents(texts))
    dimension = embeddings.shape[1]

    index = faiss.IndexFlatL2(dimension)
    index.add(embeddings)
    docstore = InMemoryDocstore({i: Document(page_content=texts[i]) for i in range(len(texts))})
    index_to_docstore_id = {i: i for i in range(len(texts))}

    vector_store = FAISS(index=index, docstore=docstore, index_to_docstore_id=index_to_docstore_id, embedding_function=embedding_model.embed_query)
    return index, docstore, index_to_docstore_id, vector_store

def preprocess_qdrant(text, embeddings, client_url, client_api_key, collection_name, batch_size=250):
    client = QdrantClient(url=client_url, api_key=client_api_key)
    client.recreate_collection(
        collection_name=collection_name,
        vectors_config=models.VectorParams(size=embeddings.shape[1], distance=models.Distance.COSINE)
    )
    
    for i in range(0, len(text), batch_size):
        batch_of_embs = embeddings[i: i + batch_size]
        batch_of_payloads = [{'page_content': doc.page_content, 'metadata': doc.metadata} for doc in text[i: i + batch_size]]
        client.upsert(collection_name=collection_name, points=models.Batch(ids=list(range(i, i + batch_size)), vectors=batch_of_embs.tolist(), payloads=batch_of_payloads))
    
    return client

def preprocess_pinecone(text, embedding_model):
    texts = [doc.page_content for doc in text]
    embeddings = np.array(embedding_model.embed_documents(texts)).tolist()
    
    pinecone_api_key = os.getenv("PINECONE_API_KEY")
    pinecone = Pinecone(api_key=pinecone_api_key, environment="us-east-1")
    index_name = "test5"
    
    if index_name in pinecone.list_indexes().names():
        pinecone.delete_index(index_name)
    
    pinecone.create_index(
        name=index_name,
        dimension=len(embeddings[0]),
        metric="cosine",
        spec=ServerlessSpec(cloud="aws", region="us-east-1"),
        deletion_protection="disabled"
    )
    
    pinecone_index = pinecone.Index(index_name)
    for i in range(0, len(texts), 100):
        batch = [(str(uuid.uuid4()), embeddings[j], {"text": texts[j]}) for j in range(i, min(i + 100, len(texts)))]
        pinecone_index.upsert(vectors=batch)
    
    return pinecone_index

def preprocess_weaviate(text, embedding_model):
    # Retrieve Weaviate credentials from environment variables
    weaviate_url = os.getenv("WEAVIATE_CLUSTER_URL")
    weaviate_api_key = os.getenv("WEAVIATE_API_KEY")
  
    # Ensure API key and URL are provided
    if not weaviate_url:
        raise ValueError("Weaviate Cluster URL is missing! Set WEAVIATE_CLUSTER_URL in environment variables.")
    if not weaviate_api_key:
        raise ValueError("Weaviate API Key is missing! Set WEAVIATE_API_KEY in environment variables.")

    try:
        # Connect to Weaviate Cloud
        client = weaviate.Client(
            url=weaviate_url,
            auth_client_secret=AuthApiKey(weaviate_api_key)
        )

        # Verify connection
        if not client.is_ready():
            raise ConnectionError("Failed to connect to Weaviate. Please check your credentials and cluster status.")

        # Store documents in Weaviate
        vs = WeaviateVectorStore.from_documents(documents=text, embedding=embedding_model, client=client)
        print("Weaviate preprocessing complete!")
        return vs

    except Exception as e:
        print(f"An error occurred while connecting to Weaviate: {e}")
        return None

def preprocess_vectordbs(files, embedding_model_name, size, overlap):
    global embedding_model_global
    embedding_model_global = SentenceTransformerEmbeddings(model_name=embedding_model_name)
    
    text = preprocess_text(files, [], size, overlap)
    st.success("Preprocessing Text Complete!")
    persist_directory = 'db'
    
    vectordb, retriever = preprocess_chroma(text, embedding_model_global, persist_directory)
    st.success("Preprocessing Chroma Complete!")
    
    index, docstore, index_to_docstore_id, vector_store = preprocess_faiss(text, embedding_model_global)
    st.success("Preprocessing Faiss Complete!")
    
    vs = preprocess_weaviate(text, embedding_model_global)
    st.success("Preprocessing Weaviate Complete!")
    
    pinecone_index = preprocess_pinecone(text, embedding_model_global)
    st.success("Preprocessing Pinecone Complete!")
    
    return index, docstore, index_to_docstore_id, vector_store, retriever, pinecone_index, embedding_model_global, vs
